import docx
import logging
from pathlib import Path
from app.ai.utils.limits import enforce_archive
from app.parser.base import BaseParser
from app.parser.exceptions import ParserError

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):
    """
    Parser implementation for DOCX resumes using python-docx.
    """
    def parse(self, file_path: Path) -> tuple[str, int]:
        logger.info(f"File detected: {file_path}")
        
        if not file_path.exists():
            # Path to the log, not to the client — see the note in pdf.py (A6).
            logger.error(f"Extraction failure: DOCX file does not exist at: {file_path}")
            raise ParserError("The uploaded file could not be read.")

        # A DOCX is a ZIP archive, and python-docx decompresses all of it before
        # returning. Every check upstream — extension, magic bytes, the 10MB
        # upload cap — measures the file on disk, so a small archive that expands
        # to gigabytes passes all of them and exhausts memory here (A3).
        #
        # Deliberately OUTSIDE the try/except below: that block converts anything
        # it catches into a generic "Failed to parse DOCX document", and a refusal
        # that names the size is worth more to the operator reading the log than
        # one that says the file was unparseable.
        enforce_archive(file_path, name=file_path.name)

        try:
            # Selected DOCX parser
            logger.info("Parser selected: DocxParser")

            # Open document
            doc = docx.Document(file_path)
            
            text_parts = []
            
            # Extract paragraph texts
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Extract table texts
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            raw_text = "\n".join(text_parts).strip()
            
            # Estimate pages: DOCX does not natively store page count.
            # 1 page ≈ 3000 characters (including spaces) is a standard estimation for resumes.
            page_count = max(1, len(raw_text) // 3000)
            
            logger.info(f"Extraction success: {file_path} (Estimated Pages: {page_count}, Bytes: {len(raw_text)})")
            return raw_text, page_count
            
        except Exception as e:
            logger.error(
                f"Extraction failure: Failed to parse DOCX {file_path}. Reason: {str(e)}",
                exc_info=True,
            )
            raise ParserError(
                "Failed to parse DOCX document. The file may be corrupt, "
                "password-protected, or not a real Word document."
            )
