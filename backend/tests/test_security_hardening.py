"""
Security hardening from the pre-launch audit: A1, A2, A3, A5, A6, A9.

Each class states the defect it closes, because "why is this asserted" is the
question a future reader has, and the answer is rarely in the assertion.

The A1 test runs a real subprocess with ENVIRONMENT=production. That is
deliberate: the docs flag is read at import time, so an in-process monkeypatch
would assert against a module that had already made its decision — a test that
passes without exercising the thing it names.
"""

from __future__ import annotations

import os
import subprocess
import sys
import zipfile
from pathlib import Path

import docx
import pytest

from app.ai.utils.limits import ArchiveTooLargeError, enforce_archive, limits
from app.core.config import settings
from app.core.observability import RateLimitMiddleware, SecurityHeadersMiddleware
from app.parser.exceptions import ParserError
from app.parser.factory import ParserFactory

BACKEND_DIR = Path(__file__).resolve().parents[1]


# ── A1 — the API documentation surface ───────────────────────────────────────
class TestDocsAreClosedInProduction:
    """`/docs`, `/redoc` and `/openapi.json` published every route and schema.

    `docs_url=None` was already set and was not enough: a custom `/docs` handler
    re-served Swagger unconditionally, and `redoc_url`/`openapi_url` were never
    set at all, so they kept their framework defaults and answered 200.
    """

    _PROBE = (
        "import os, json\n"
        "from fastapi.testclient import TestClient\n"
        "from app.main import app, DOCS_ENABLED\n"
        "c = TestClient(app)\n"
        "print(json.dumps({\n"
        "  'enabled': DOCS_ENABLED,\n"
        "  'docs': c.get('/docs').status_code,\n"
        "  'redoc': c.get('/redoc').status_code,\n"
        "  'openapi': c.get('/openapi.json').status_code,\n"
        "  'health': c.get('/health').status_code,\n"
        "}))\n"
    )

    def _probe(self, environment: str) -> dict:
        env = {
            **os.environ,
            "ENVIRONMENT": environment,
            # Production refuses to boot without explicit origins (A2). Supplied
            # so this test measures the docs surface and not the CORS guard.
            "ALLOWED_ORIGINS": "https://app.example.com",
        }
        out = subprocess.run(
            [sys.executable, "-c", self._PROBE],
            cwd=BACKEND_DIR, env=env, capture_output=True, text=True, timeout=180,
        )
        assert out.returncode == 0, out.stderr[-2000:]
        import json
        return json.loads(out.stdout.strip().splitlines()[-1])

    def test_production_serves_no_docs(self):
        r = self._probe("production")
        assert r["enabled"] is False
        assert r["docs"] == 404, "Swagger UI is reachable in production"
        assert r["redoc"] == 404, "ReDoc is reachable in production"
        assert r["openapi"] == 404, "the OpenAPI schema is downloadable in production"

    def test_production_still_serves_the_application(self):
        """A closed docs surface must not be a broken service."""
        assert self._probe("production")["health"] == 200

    def test_development_is_unchanged(self):
        """The whole point of the flag: local behaviour is exactly as before."""
        r = self._probe("development")
        assert r["enabled"] is True
        assert (r["docs"], r["redoc"], r["openapi"]) == (200, 200, 200)


# ── A2 — CORS fails closed ───────────────────────────────────────────────────
class TestCorsFailsClosedInProduction:
    """The default was `"*"`, so forgetting the variable served every origin."""

    def test_the_insecure_default_is_gone(self):
        from app.core.config import Settings

        assert Settings.model_fields["ALLOWED_ORIGINS"].default == ""

    def test_unset_in_production_yields_no_origins(self, monkeypatch):
        monkeypatch.setattr(settings, "ALLOWED_ORIGINS", "")
        monkeypatch.setattr(settings, "ENVIRONMENT", "production")
        assert settings.allowed_origins == []

    def test_unset_in_development_stays_convenient(self, monkeypatch):
        """Developers keep the wildcard; only production is strict."""
        monkeypatch.setattr(settings, "ALLOWED_ORIGINS", "")
        monkeypatch.setattr(settings, "ENVIRONMENT", "development")
        assert settings.allowed_origins == ["*"]

    def test_explicit_origins_always_win(self, monkeypatch):
        monkeypatch.setattr(settings, "ALLOWED_ORIGINS", "https://a.com, https://b.com")
        monkeypatch.setattr(settings, "ENVIRONMENT", "production")
        assert settings.allowed_origins == ["https://a.com", "https://b.com"]

    def test_startup_refuses_unset_origins_in_production(self, monkeypatch):
        """The case that used to slip through — absent, not wildcard."""
        from app.core.startup import validate_startup

        monkeypatch.setattr(settings, "ALLOWED_ORIGINS", "")
        monkeypatch.setattr(settings, "ENVIRONMENT", "production")
        with pytest.raises(RuntimeError, match="ALLOWED_ORIGINS is not set"):
            validate_startup()

    def test_startup_still_refuses_the_wildcard_in_production(self, monkeypatch):
        from app.core.startup import validate_startup

        monkeypatch.setattr(settings, "ALLOWED_ORIGINS", "*")
        monkeypatch.setattr(settings, "ENVIRONMENT", "production")
        with pytest.raises(RuntimeError, match="allow all origins"):
            validate_startup()


# ── A3 — DOCX decompression bomb ─────────────────────────────────────────────
def _docx_with(tmp_path: Path, paragraphs: int, words: int = 400) -> Path:
    path = tmp_path / "sample.docx"
    document = docx.Document()
    line = "COMPRESSIBLE " * words
    for _ in range(paragraphs):
        document.add_paragraph(line)
    document.save(path)
    return path


class TestDocxArchiveIsBounded:
    """A DOCX is a ZIP. Extension, magic bytes and the 10MB upload cap all
    measure the file on DISK; python-docx decompresses all of it into memory.
    Measured: ordinary prose reaches 300:1, so 10MB expands to roughly 3GB."""

    def test_a_normal_resume_still_parses(self, tmp_path):
        path = _docx_with(tmp_path, paragraphs=3, words=5)
        text, pages, parser = ParserFactory.parse_file(path)
        assert parser == "DocxParser" and pages >= 1 and "COMPRESSIBLE" in text

    def test_an_oversized_archive_is_refused(self, tmp_path, monkeypatch):
        monkeypatch.setattr(settings, "AI_MAX_ARCHIVE_UNCOMPRESSED_MB", 1)
        path = _docx_with(tmp_path, paragraphs=2000)
        with pytest.raises(ArchiveTooLargeError):
            enforce_archive(path, name=path.name)

    def test_the_refusal_reaches_routes_as_a_parser_error(self, tmp_path, monkeypatch):
        """Routes already handle ParserError as a 400. An unrecognised exception
        type here would surface as a 500 — the wrong answer for a bad upload."""
        monkeypatch.setattr(settings, "AI_MAX_ARCHIVE_UNCOMPRESSED_MB", 1)
        path = _docx_with(tmp_path, paragraphs=2000)
        with pytest.raises(ParserError):
            ParserFactory.parse_file(path)

    def test_the_message_names_both_numbers(self, tmp_path, monkeypatch):
        monkeypatch.setattr(settings, "AI_MAX_ARCHIVE_UNCOMPRESSED_MB", 1)
        path = _docx_with(tmp_path, paragraphs=2000)
        with pytest.raises(ArchiveTooLargeError) as excinfo:
            enforce_archive(path, name=path.name)
        assert "above the limit of 1MB" in str(excinfo.value)

    def test_the_guard_runs_before_python_docx_opens_the_file(self, tmp_path, monkeypatch):
        """The ordering IS the fix. `enforce_document` bounds extracted text, but
        by the time it runs the archive is already decompressed and the memory is
        already spent."""
        monkeypatch.setattr(settings, "AI_MAX_ARCHIVE_UNCOMPRESSED_MB", 1)
        # Build the fixture BEFORE patching: `_docx_with` calls docx.Document()
        # itself to author the file, and a patch installed earlier records that
        # call instead of the parser's — which is what the first version of this
        # test did, failing against its own helper rather than the guard.
        path = _docx_with(tmp_path, paragraphs=2000)

        opened: list[Path] = []
        real = docx.Document
        monkeypatch.setattr(
            docx, "Document", lambda p=None, *a, **k: (opened.append(p), real(p))[1]
        )
        with pytest.raises(ParserError):
            ParserFactory.parse_file(path)
        assert opened == [], "python-docx opened the archive despite the guard"

    def test_a_non_zip_is_left_to_the_parser(self, tmp_path):
        """Size is this function's job; diagnosing corruption is not."""
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"PK\x03\x04 definitely not a real archive")
        enforce_archive(path, name=path.name)  # must not raise

    def test_the_ceiling_is_configurable_and_owned_by_the_limits_module(self):
        assert limits().archive_uncompressed_bytes == (
            settings.AI_MAX_ARCHIVE_UNCOMPRESSED_MB * 1024 * 1024
        )

    def test_the_default_ceiling_admits_any_real_resume(self):
        assert limits().archive_uncompressed_bytes >= 50 * 1024 * 1024


# ── A5 — rate limiting reaches the authenticated AI endpoints ────────────────
class TestAuthenticatedAiEndpointsAreRateLimited:
    """Being authenticated bounds WHO can spend, not how fast. Every path below
    is a paid Groq call and none of them was limited."""

    def _rule(self, path):
        return RateLimitMiddleware(app=None)._rule(path)

    @pytest.mark.parametrize("path", [
        "/api/v1/campaigns/abc123/compare",
        "/api/v1/campaigns/abc123/candidates/def456/interview",
        "/api/v1/campaigns/abc123/candidates/def456/resume",
        "/api/v1/campaigns/abc123/embeddings/reindex",
        "/api/v1/agent/scan",
    ])
    def test_expensive_authenticated_paths_have_a_rule(self, path):
        assert self._rule(path) is not None, f"{path} is unlimited"

    def test_the_previously_covered_paths_are_unchanged(self):
        assert self._rule("/api/v1/batch-analysis") == (20, 60)
        assert self._rule("/api/v1/copilot/chat") == (30, 60)

    def test_cheap_reads_are_still_unlimited(self):
        """The limiter is for spend, not for traffic."""
        for path in ("/api/v1/campaigns", "/health", "/api/v1/activity"):
            assert self._rule(path) is None

    def test_the_resume_url_read_is_not_caught_by_the_resume_rule(self):
        """`/resume-url` is a GET that issues a signed link — a suffix rule that
        swallowed it would rate-limit ordinary page loads."""
        assert self._rule("/api/v1/campaigns/a/candidates/b/resume-url") is None


# ── A6 — no raw filesystem exceptions in responses ───────────────────────────
class TestUploadErrorsDoNotLeakInternals:
    def test_the_source_returns_no_exception_text(self):
        """Both 500 paths interpolated `str(e)` — absolute paths and OS
        permission detail, useful in a log, free recon in a response body."""
        source = (BACKEND_DIR / "app" / "services" / "upload_utils.py").read_text(encoding="utf-8")
        assert "detail=f\"Error saving file: {str(e)}\"" not in source
        assert "detail=f\"Could not create upload directory: {str(e)}\"" not in source
        assert source.count("Could not process the upload") == 2

    def test_the_exception_is_still_logged(self):
        source = (BACKEND_DIR / "app" / "services" / "upload_utils.py").read_text(encoding="utf-8")
        assert source.count("logger.exception") >= 2, "swallowed without a log"


# ── A9 — authenticated responses are not cacheable ───────────────────────────
class TestApiResponsesAreNotStored:
    def test_no_store_is_applied(self):
        from fastapi.testclient import TestClient
        from app.main import app

        assert TestClient(app).get("/health").headers.get("Cache-Control") == "no-store"

    def test_it_is_a_default_a_route_can_override(self):
        """`setdefault`, like every other header here — a route that deliberately
        sets caching keeps it, which is what makes this safe to apply blanket."""
        source = (BACKEND_DIR / "app" / "core" / "observability.py").read_text(encoding="utf-8")
        assert 'setdefault("Cache-Control", "no-store")' in source
